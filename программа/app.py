import sqlite3
from flask import Flask, render_template, redirect, request, flash, send_from_directory
from werkzeug.exceptions import abort
import os
import pprint
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = b'my)secret)key'
UPLOAD_FOLDER = 'contracts'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def get_db_connection():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    return conn

@app.route('/')
def index():
    return redirect("/contracts")

###############
###############
############### Контракты 
###############
###############

@app.route('/contracts')
def contracts():
    """ Страница-список - получение всех контрактов """

    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, events, clients
        WHERE contracts.event_id = events.id_event and contracts.client_id = clients.id_client
    """).fetchall()
    conn.close()
    return render_template('contracts.html', contracts=pos)


def get_contract(item_id):
    """ Получение одного контракта из БД """
    
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM contracts, events, clients
        WHERE contracts.event_id = events.id_event and contracts.client_id = clients.id_client and contracts.id_contract = ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/contract/<int:contract_id>')
def contract(contract_id):
    """ Страница-карточка - 1 контракт """

    pos = get_contract(contract_id)
    return render_template('contract.html', contract=pos)


@app.route('/new_contract', methods=('GET', 'POST'))
def new_contract():
    """ Страница-добавления нового контракта """

    if request.method == 'POST':
        # добавление нового контракта в БД псоле заполнения формы
        try:
            number = request.form['number']
            date = request.form['date']
            price = int(request.form['price'])
            deal_status = 0
            client_id = int(request.form.get('owner'))
            event_id = int(request.form.get('event'))
            employee_id = int(request.form.get('employee'))
        except ValueError:
            flash('Некорректные значения')
            client_id = 0
        if not (client_id > 0 and event_id > 0 and employee_id > 0):
            flash('Не все поля заполнены')
        else:
            if not (number and date and price):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute("INSERT INTO 'contracts' ('number', 'date', 'price', 'deal_status', 'event_id', 'client_id', 'employee_id')  VALUES (?, ?, ?, ?, ?, ?, ?)",
                    (number, date, price, deal_status, event_id, client_id, employee_id))
                conn.commit()
                new_contract_id = cursor.lastrowid
                print(new_contract_id)
                conn.close()
                return redirect(f'/contract/{new_contract_id}')
    
    # отрисовка формы        
    conn = get_db_connection()
    pos1 = conn.execute("""SELECT * FROM clients""").fetchall()
    pos2 = conn.execute("""SELECT * FROM events""").fetchall()
    pos3 = conn.execute("""SELECT * FROM employees""").fetchall()
    conn.close()
    return render_template('new_contract.html', clients=pos1, events=pos2, employees=pos3)


@app.route('/generate_contract', methods=('GET', 'POST'))
def generate_contract():
    """ Страница генерации договора """

    # переменные шаблона
    id = int(request.args.get('id_contract'))
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, events, clients, employees
            WHERE contracts.event_id = events.id_event and
                        contracts.client_id = clients.id_client and 
                        contracts.employee_id = employees.id_employee and
                        contracts.id_contract = ?
                        """, (id,)).fetchone()
    conn.close()
    contract_params = {
            'VERNEL _NEED': 'Средство для белья ',
            'CIF_NEED': 'Средство для пола ',
            'SANOKS_NEED': 'Средство для стекол ',
            'VANISH_NEED': 'Средство для ванной комнаты ',
            'SPIRT_NEED': 'Средство для обработки различных поверхностей ',
            'OTHER_INFO': 'Дополнительные требования клиента'}
    contract_params_auto = {
            'CONTRACT_NUMBER': ['номер договора', pos['number']],
            'CONTRACT_DATE': ['дата подписания договора', pos['date']],
            'EMPLOYEE_FULLNAME': ['ФИО сотрудника ', pos['name']],
            'HOTEL_ADDRESS': ['адрес', pos['address']],
            'HOTEL_DATE': ['дата', pos['date']],}

    if request.method == 'POST':
        # создание нового документа
        result_params =  request.form.to_dict()
        create_contract(id, result_params)
        return redirect(f'/contract/{id}')

    # скачивание файла, если он заполнен
    filename = f"Отчет уборка {pos['number']} от {pos['date']}.docx"
    if os.path.exists(os.path.join('contracts', filename)):
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)
    else:
        # отрисовка формы заполнения
        flash('Отчет не сформирован, заполните его')
        return render_template('generate_contract.html', 
                               contract=pos, contract_params=contract_params, auto_params=contract_params_auto)


def create_contract(id, contract_params):
    """ Создание нового документа по шаблону """

    template = os.path.join('contracts', 'contract_template.docx')
    result = os.path.join('contracts', f"Отчет уборка {contract_params['CONTRACT_NUMBER']} от {contract_params['CONTRACT_DATE']}.docx")
    
    template_doc = Document(template)
    for key, value in contract_params.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, f'=={key}==', value)
        for table in template_doc.tables:
            replace_text_in_tables(table, key, value)
    template_doc.save(result)


def replace_text(paragraph, key, value):
  """ Работа docx - заполнение параграфов """

  if key in paragraph.text:
    paragraph.text = paragraph.text.replace(key, value)


def replace_text_in_tables(table, key, value):
  """ Работа docx - заполнение таблиц """

  for row in table.rows:
    for cell in row.cells:
      if key in cell.text:
        cell.text = cell.text.replace(key, value)

###############
###############
############### Мероприятия
###############
###############

@app.route('/events')
def events():
    """ Страница-список - получение всех мероприятий """

    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, events, clients
    WHERE contracts.event_id = events.id_event and contracts.client_id = clients.id_client
    """).fetchall()
    conn.close()
    return render_template('hotels.html', events=pos)


def get_event(item_id):
    """ Получение одного мероприятия из БД """

    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM contracts, events, clients
        WHERE contracts.event_id = events.id_event and contracts.client_id = clients.id_client and events.id_event = ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/event/<int:event_id>')
def event(event_id):
    """ Страница-карточка - 1 мероприятие """

    pos = get_event(event_id)
    return render_template('hotel.html', event=pos)


@app.route('/new_event', methods=('GET', 'POST'))
def new_event():
    """ Страница-добавления нового мероприятия """

    if request.method == 'POST':
        # добавление нового мероприятия в БД псоле заполнения формы
        address = ", ".join([request.form['region'],
                                request.form['city'],
                                "Ул. " + request.form['street'],
                                "Номер " + request.form['house'],
        ]).lstrip(", ")
        try:
            people = float(request.form['people'])
            date = float(request.form['date'])
            description = request.form['description']
            client_id = int(request.form.get('owner'))
        except ValueError:
            flash('Некорректные значения')
            client_id = 0
        if not client_id > 0:
            flash('Не все поля заполнены')
        else:
            if not (address and people):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                conn.execute("INSERT INTO 'events' ('address', 'people', 'date', 'description', 'client_id')  VALUES (?, ?, ?, ?, ?)",
                    (address, people, date, description, client_id))
                conn.commit()
                conn.close()
                return redirect(f'/new_contract')

    # отрисовка формы
    conn = get_db_connection()
    pos1 = conn.execute("""SELECT * FROM clients""").fetchall()
    pos2 = conn.execute("""SELECT * FROM events""").fetchall()
    conn.close()
    return render_template('new_hotel.html', clients=pos1, events=pos2)

#
# Клиенты 
#


@app.route('/clients')
def clients():
    """ Страница-список - получение всех мероприятий """

    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, events, clients
    WHERE contracts.event_id = events.id_event and contracts.client_id = clients.id_client
    """).fetchall()
    conn.close()
    return render_template('clients.html', clients=pos)


def get_client(client_id):
    """ Получение одного мероприятия из БД """

    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM clients
        WHERE id_client = ?""", (client_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/client/<int:client_id>')
def client(client_id):
    """ Страница-карточка - 1 мероприятие """

    pos = get_client(client_id)
    return render_template('client.html', client=pos)

#
# Отчеты 
#

@app.route('/reports')
def reports():
    abort(404)

#
# Сотрудники 
#

@app.route('/employees')
def employees():
    abort(404)

#
# 404 
#

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404


if __name__ == '__main__':
    app.run()